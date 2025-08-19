import os
from flask import Flask, render_template, request
from docx import Document  
from docx.shared import Pt
import io
import PyPDF2
import joblib
import pandas as pd



app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024


model = None
model_columns = None
if os.path.exists('document_classifier.joblib') and os.path.exists('model_columns.joblib'):
    try:
        model = joblib.load('document_classifier.joblib')
        model_columns = joblib.load('model_columns.joblib')
    except Exception as e:
        print("Gagal memuat model ML:", e)

def check_docx(doc: Document, custom_settings=None) -> dict:
    report = []
    success = True
    

    settings = custom_settings or {
        'marginLeft': 4.0,
        'marginRight': 3.0,
        'marginTop': 3.0,
        'marginBottom': 3.0,
        'lineSpacing': 1.5,
        'fontSize': 12,
        'marginTolerance': 0.1
    }


    if settings.get('enableFontTypeCheck', False):
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name != 'Times New Roman':
                    report.append(f'Font tidak sesuai di paragraf: "{para.text[:30]}..."')
                    success = False
                    break


    if settings.get('enableFontSizeCheck', True):
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size and run.font.size.pt != settings['fontSize']:
                    report.append(f'Ukuran font tidak sesuai di paragraf: "{para.text[:30]}..." (Diharapkan: {settings["fontSize"]}pt)')
                    success = False
                    break


    if settings.get('enableSpacingCheck', True):
        for para in doc.paragraphs:
            if para.paragraph_format.line_spacing is not None:
                spacing = para.paragraph_format.line_spacing
                if abs(spacing - settings['lineSpacing']) > 0.1:
                    report.append(f'Spasi tidak sesuai di paragraf: "{para.text[:30]}..." (Diharapkan: {settings["lineSpacing"]})')
                    success = False


    if settings.get('enableMarginCheck', True):
        sections = doc.sections
        if sections:
            section = sections[0]

            left_margin_cm = section.left_margin.inches * 2.54
            right_margin_cm = section.right_margin.inches * 2.54
            top_margin_cm = section.top_margin.inches * 2.54
            bottom_margin_cm = section.bottom_margin.inches * 2.54

            expected_margins_cm = {
                'left': settings['marginLeft'],
                'right': settings['marginRight'],
                'top': settings['marginTop'],
                'bottom': settings['marginBottom']
            }

            tolerance = settings['marginTolerance']

            if abs(left_margin_cm - expected_margins_cm['left']) > tolerance:
                report.append(f'Margin kiri tidak sesuai: {left_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["left"]} cm)')
                success = False
            if abs(right_margin_cm - expected_margins_cm['right']) > tolerance:
                report.append(f'Margin kanan tidak sesuai: {right_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["right"]} cm)')
                success = False
            if abs(top_margin_cm - expected_margins_cm['top']) > tolerance:
                report.append(f'Margin atas tidak sesuai: {top_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["top"]} cm)')
                success = False
            if abs(bottom_margin_cm - expected_margins_cm['bottom']) > tolerance:
                report.append(f'Margin bawah tidak sesuai: {bottom_margin_cm:.2f} cm (Diharapkan: {expected_margins_cm["bottom"]} cm)')
                success = False
        else:
            if settings.get('enableMarginCheck', True):
                report.append('Tidak dapat memeriksa margin.')
                success = False

    result = {"success": success, "messages": report}


    if model and model_columns:
        try:
            features = extract_features_docx(doc)
            features_df = prepare_features(features)
            prediction = model.predict(features_df)[0]
            prediction_prob = model.predict_proba(features_df)[0][1]

            ml_message = "Dokumen sesuai dengan kriteria (Prediksi ML: Correct)." if prediction == 1 else "Dokumen tidak sesuai dengan kriteria (Prediksi ML: Incorrect)."
            result["messages"].append(ml_message)
            result["confidence"] = f"{prediction_prob * 100:.2f}%"
            result["success"] = result["success"] and (prediction == 1)
        except Exception as e:
            result["messages"].append("Gagal melakukan prediksi ML pada dokumen.")
            result["success"] = False

    return result

def check_pdf(file_stream: io.BytesIO, custom_settings=None) -> dict:
    report = []
    success = True
    try:
        file_stream.seek(0)
        pdf_reader = PyPDF2.PdfReader(file_stream)
        

        if len(pdf_reader.pages) == 0:
            report.append('PDF tidak memiliki halaman yang dapat dibaca.')
            return {"success": False, "messages": report}
            

        report.append('PDF berhasil dibaca dan divalidasi.')
        report.append('Catatan: Pemeriksaan detail font dan format pada PDF memiliki keterbatasan dengan library saat ini.')
        report.append('Disarankan menggunakan format .docx untuk pemeriksaan yang lebih akurat.')
        

        success = True
        
    except Exception as e:
        report.append(f'Gagal membaca dokumen PDF: {str(e)}')
        return {"success": False, "messages": report}

    result = {"success": success, "messages": report}

    if model and model_columns:
        try:
            features = extract_features_pdf(file_stream)
            features_df = prepare_features(features)
            prediction = model.predict(features_df)[0]
            prediction_prob = model.predict_proba(features_df)[0][1]

            ml_message = "Dokumen sesuai dengan kriteria (Prediksi ML: Correct)." if prediction == 1 else "Dokumen tidak sesuai dengan kriteria (Prediksi ML: Incorrect)."
            result["messages"].append(ml_message)
            result["confidence"] = f"{prediction_prob * 100:.2f}%"
            result["success"] = result["success"] and (prediction == 1)
        except Exception as e:
            result["messages"].append("Gagal melakukan prediksi ML pada dokumen PDF.")
            result["success"] = False

    return result

def extract_features_docx(doc: Document) -> dict:
    features = {}
    fonts = {}
    sizes = {}
    line_spacings = set()
    margins = {}

    for para in doc.paragraphs:
        for run in para.runs:
            font = run.font.name or "Unknown"
            size = run.font.size.pt if run.font.size else 12
            fonts[font] = fonts.get(font, 0) + 1
            sizes[size] = sizes.get(size, 0) + 1
        if para.paragraph_format.line_spacing:
            line_spacings.add(para.paragraph_format.line_spacing)


    if doc.sections:
        section = doc.sections[0]
        margins['left_margin_cm'] = section.left_margin.inches * 2.54
        margins['right_margin_cm'] = section.right_margin.inches * 2.54
        margins['top_margin_cm'] = section.top_margin.inches * 2.54
        margins['bottom_margin_cm'] = section.bottom_margin.inches * 2.54
    else:
        margins = {'left_margin_cm': 0, 'right_margin_cm': 0, 'top_margin_cm': 0, 'bottom_margin_cm': 0}


    features['unique_fonts'] = len(fonts)
    features['most_common_font'] = max(fonts, key=fonts.get) if fonts else "Unknown"
    features['font_size_variance'] = pd.Series(list(sizes.keys())).std() if sizes else 0
    features['unique_line_spacings'] = len(line_spacings)

    for key, value in margins.items():
        features[key] = value

    return features

def extract_features_pdf(file_stream: io.BytesIO) -> dict:
    features = {}
    try:
        file_stream.seek(0)
        pdf_reader = PyPDF2.PdfReader(file_stream)
        
        features['unique_fonts'] = 1
        features['most_common_font'] = "Unknown"
        features['font_size_variance'] = 0
        features['left_margin_cm'] = 0
        features['right_margin_cm'] = 0
        features['top_margin_cm'] = 0
        features['bottom_margin_cm'] = 0
        
    except Exception as e:
        print("Gagal membuka PDF:", e)

        features = {
            'unique_fonts': 1,
            'most_common_font': "Unknown",
            'font_size_variance': 0,
            'left_margin_cm': 0,
            'right_margin_cm': 0,
            'top_margin_cm': 0,
            'bottom_margin_cm': 0
        }

    return features

def prepare_features(features: dict) -> pd.DataFrame:
    df = pd.DataFrame([features])
    df = pd.get_dummies(df, columns=['most_common_font'], drop_first=True)

    for col in model_columns:
        if col not in df.columns:
            df[col] = 0

    df = df[model_columns]
    df.fillna(0, inplace=True)

    return df

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        uploaded_files = request.files.getlist('files')
        if not uploaded_files or uploaded_files == [None]:
            return render_template('index.html', reports=[{"success": False, "filename": "Tidak ada file yang diunggah.", "messages": ["Tidak ada file yang diunggah."]}])


        custom_settings = {}
        

        for key in ['marginLeft', 'marginRight', 'marginTop', 'marginBottom', 'lineSpacing', 'fontSize', 'marginTolerance']:
            setting_key = f'setting_{key}'
            if setting_key in request.form:
                try:
                    if key == 'fontSize':
                        custom_settings[key] = int(float(request.form[setting_key]))
                    else:
                        custom_settings[key] = float(request.form[setting_key])
                except ValueError:
                    pass
        

        for key in ['enableMarginCheck', 'enableSpacingCheck', 'enableFontSizeCheck', 'enableFontTypeCheck']:
            setting_key = f'setting_{key}'
            if setting_key in request.form:
                custom_settings[key] = request.form[setting_key].lower() == 'true'

        reports = []
        for uploaded_file in uploaded_files:
            if uploaded_file.filename == '':
                continue

            file_name = uploaded_file.filename.lower()
            if not (file_name.endswith('.docx') or file_name.endswith('.pdf')):
                reports.append({
                    "success": False,
                    "filename": uploaded_file.filename,
                    "messages": ["Silakan unggah file .docx atau .pdf saja."]
                })
                continue

            file_bytes = uploaded_file.read()
            file_stream = io.BytesIO(file_bytes)

            if file_name.endswith('.docx'):
                try:
                    doc = Document(file_stream)
                    report = check_docx(doc, custom_settings if custom_settings else None)
                except Exception as e:
                    report = {"success": False, "messages": ["Gagal membaca dokumen .docx. Pastikan file dalam format yang benar."]}
            elif file_name.endswith('.pdf'):
                report = check_pdf(file_stream, custom_settings if custom_settings else None)

            report["filename"] = uploaded_file.filename
            reports.append(report)

        return render_template('index.html', reports=reports)

    return render_template('index.html')

@app.errorhandler(413)
def request_entity_too_large(error):
    return render_template('index.html', reports=[{"success": False, "filename": "File terlalu besar.", "messages": ["File terlalu besar. Maksimal 50MB."]}]), 413

if __name__ == '__main__':
    print("Isi folder 'templates':", os.listdir('templates'))
    app.run(host='0.0.0.0', port=int(os.environ.get("PORT", 81)))
